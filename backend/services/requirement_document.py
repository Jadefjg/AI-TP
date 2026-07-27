from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import httpx

from backend.core.config import get_settings

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx", ".pdf", ".html", ".htm"}
MAX_FETCH_BYTES = 10 * 1024 * 1024
MIN_REQUIREMENT_CHARS = 10

_BROWSER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,application/json,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}

_FEISHU_HOST_RE = re.compile(r"(?:^|\.)(?:feishu\.cn|larksuite\.com|larkoffice\.com)$", re.I)
_FEISHU_PATH_RE = re.compile(
    r"/(?P<kind>wiki|docx|docs|doc)/(?P<token>[A-Za-z0-9_-]{8,})",
    re.I,
)
_LOGIN_MARKERS = (
    "suite-passport",
    "login required",
    "passport_web_did",
    "accounts.feishu",
    "accounts.larksuite",
)


def parse_requirement_document(*, filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {suffix or 'unknown'}，支持 {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    if suffix in {".txt", ".md", ".markdown"}:
        text = _decode_text(content)
        if len(text.strip()) < MIN_REQUIREMENT_CHARS:
            raise ValueError(f"文档内容过短，至少需要 {MIN_REQUIREMENT_CHARS} 个字符")
        return text.strip(), suffix.lstrip(".")

    if suffix == ".docx":
        return _parse_docx(content), "docx"

    if suffix == ".pdf":
        return _parse_pdf(content), "pdf"

    if suffix in {".html", ".htm"}:
        return _parse_html(content), "html"

    raise ValueError("无法解析文档")


def fetch_requirement_from_url(url: str) -> tuple[str, str, str]:
    """Download a document link and extract requirement text. Returns (text, format, source_label)."""
    raw_url = url.strip()
    if not raw_url.startswith(("http://", "https://")):
        raise ValueError("仅支持 http/https 文档链接")

    feishu = _parse_feishu_doc_url(raw_url)
    if feishu:
        return _fetch_feishu_document(raw_url, kind=feishu[0], token=feishu[1])

    with httpx.Client(timeout=30.0, follow_redirects=True, headers=_BROWSER_HEADERS) as client:
        response = client.get(raw_url)
        response.raise_for_status()
        content = response.content
        if len(content) > MAX_FETCH_BYTES:
            raise ValueError("文档过大，请上传本地文件")

    path_name = Path(urlparse(raw_url).path).name
    suffix = Path(path_name).suffix.lower()
    content_type = (response.headers.get("content-type") or "").lower()

    if suffix in SUPPORTED_EXTENSIONS or "pdf" in content_type:
        filename = path_name if suffix else "document.pdf"
        if suffix not in SUPPORTED_EXTENSIONS and "pdf" in content_type:
            filename = "document.pdf"
        text, fmt = parse_requirement_document(filename=filename, content=content)
        return text, fmt, raw_url

    if suffix == ".docx" or "wordprocessingml" in content_type:
        text, fmt = parse_requirement_document(filename=path_name or "document.docx", content=content)
        return text, fmt, raw_url

    if "html" in content_type or suffix in {".html", ".htm"} or _looks_like_html(content):
        return _parse_html(content, source_url=raw_url), "html", raw_url

    # Fallback: try plain text decode
    text = _decode_text(content).strip()
    if len(text) >= MIN_REQUIREMENT_CHARS:
        return text, "text", raw_url

    raise ValueError("无法从链接解析需求文档，请上传 Word/PDF 文件，或粘贴需求原文")


def _looks_like_html(content: bytes) -> bool:
    head = content[:2048].lstrip().lower()
    return head.startswith(b"<!doctype html") or head.startswith(b"<html")


def _parse_feishu_doc_url(url: str) -> tuple[str, str] | None:
    parsed = urlparse(url)
    host = (parsed.hostname or "").lower()
    if not host or not _FEISHU_HOST_RE.search(host):
        return None
    match = _FEISHU_PATH_RE.search(parsed.path or "")
    if not match:
        return None
    return match.group("kind").lower(), match.group("token")


def _fetch_feishu_document(url: str, *, kind: str, token: str) -> tuple[str, str, str]:
    settings = get_settings()
    app_id = (settings.feishu_app_id or "").strip()
    app_secret = (settings.feishu_app_secret or "").strip()

    if app_id and app_secret:
        try:
            text = _fetch_feishu_via_open_api(
                kind=kind,
                token=token,
                app_id=app_id,
                app_secret=app_secret,
                open_base=settings.feishu_open_base_url.strip() or "https://open.feishu.cn",
            )
            if len(text.strip()) >= MIN_REQUIREMENT_CHARS:
                return text.strip(), "feishu", url
        except ValueError:
            raise
        except Exception as exc:  # noqa: BLE001 - surface as user-facing fetch error
            raise ValueError(f"飞书 Open API 拉取失败：{exc}") from exc

    # Public/anonymous HTML fetch usually hits login SPA for private wiki pages.
    with httpx.Client(timeout=30.0, follow_redirects=True, headers=_BROWSER_HEADERS) as client:
        response = client.get(url)
        response.raise_for_status()
        content = response.content
        content_type = (response.headers.get("content-type") or "").lower()
        if "json" in content_type:
            payload = response.json()
            if isinstance(payload, dict) and int(payload.get("code") or 0) in {5, 99991663}:
                raise ValueError(_feishu_login_required_message(has_open_api=bool(app_id and app_secret)))

    raw = _decode_text(content).lower()
    if any(marker in raw for marker in _LOGIN_MARKERS):
        raise ValueError(_feishu_login_required_message(has_open_api=bool(app_id and app_secret)))

    try:
        text = _parse_html(content, source_url=url)
    except ValueError as exc:
        if "过短" in str(exc):
            raise ValueError(_feishu_login_required_message(has_open_api=bool(app_id and app_secret))) from exc
        raise
    return text, "feishu", url


def _feishu_login_required_message(*, has_open_api: bool) -> str:
    if has_open_api:
        return (
            "飞书文档需要登录，且当前应用凭证未能读取该文档。"
            "请确认：1) 文档已对应用授权（云文档「…」→ 添加文档应用）；"
            "2) 应用具备 wiki/docx 读权限；或改为导出 Word/PDF 上传 / 粘贴原文。"
        )
    return (
        "飞书/Lark 文档需登录，无法通过公开网页直接抓取正文。"
        "请任选其一：1) 将文档设为「互联网获得链接的人可阅读」并重试；"
        "2) 在 .env 配置 FEISHU_APP_ID / FEISHU_APP_SECRET 并用应用读取文档；"
        "3) 导出为 Word/PDF/Markdown 后上传，或直接粘贴需求原文。"
    )


def _fetch_feishu_via_open_api(
    *,
    kind: str,
    token: str,
    app_id: str,
    app_secret: str,
    open_base: str,
) -> str:
    base = open_base.rstrip("/")
    with httpx.Client(timeout=30.0, follow_redirects=True) as client:
        token_resp = client.post(
            f"{base}/open-apis/auth/v3/tenant_access_token/internal",
            json={"app_id": app_id, "app_secret": app_secret},
        )
        token_resp.raise_for_status()
        token_payload = token_resp.json()
        if int(token_payload.get("code") or 0) != 0:
            raise ValueError(token_payload.get("msg") or "获取 tenant_access_token 失败")
        access_token = str(token_payload.get("tenant_access_token") or "").strip()
        if not access_token:
            raise ValueError("飞书未返回 tenant_access_token")

        headers = {"Authorization": f"Bearer {access_token}"}
        obj_token = token
        obj_type = "docx" if kind in {"wiki", "docx"} else kind

        if kind == "wiki":
            node_resp = client.get(
                f"{base}/open-apis/wiki/v2/spaces/get_node",
                params={"token": token},
                headers=headers,
            )
            node_resp.raise_for_status()
            node_payload = node_resp.json()
            if int(node_payload.get("code") or 0) != 0:
                raise ValueError(node_payload.get("msg") or "获取知识库节点失败")
            node = ((node_payload.get("data") or {}).get("node") or {}) if isinstance(node_payload, dict) else {}
            obj_token = str(node.get("obj_token") or token).strip()
            obj_type = str(node.get("obj_type") or "docx").strip().lower()

        return _feishu_raw_content(client, base=base, headers=headers, obj_token=obj_token, obj_type=obj_type)


def _feishu_raw_content(
    client: httpx.Client,
    *,
    base: str,
    headers: dict[str, str],
    obj_token: str,
    obj_type: str,
) -> str:
    if obj_type in {"docx", "wiki"}:
        resp = client.get(f"{base}/open-apis/docx/v1/documents/{obj_token}/raw_content", headers=headers)
        resp.raise_for_status()
        payload = resp.json()
        if int(payload.get("code") or 0) != 0:
            raise ValueError(payload.get("msg") or "读取 docx 正文失败")
        content = ((payload.get("data") or {}).get("content") or "") if isinstance(payload, dict) else ""
        text = str(content).strip()
        if len(text) >= MIN_REQUIREMENT_CHARS:
            return text
        raise ValueError("飞书文档正文过短或为空")

    if obj_type in {"doc", "docs"}:
        resp = client.get(f"{base}/open-apis/doc/v2/{obj_token}/raw_content", headers=headers)
        resp.raise_for_status()
        payload = resp.json()
        if int(payload.get("code") or 0) != 0:
            raise ValueError(payload.get("msg") or "读取 docs 正文失败")
        content = ((payload.get("data") or {}).get("content") or "") if isinstance(payload, dict) else ""
        text = str(content).strip()
        if len(text) >= MIN_REQUIREMENT_CHARS:
            return text
        raise ValueError("飞书文档正文过短或为空")

    raise ValueError(f"暂不支持飞书文档类型：{obj_type}，请导出 Word/PDF 后上传")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文档编码")


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，无法解析 Word 文档") from exc

    doc = Document(BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        line = (para.text or "").strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").strip()) for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts).strip()
    if len(text) < MIN_REQUIREMENT_CHARS:
        raise ValueError("Word 文档解析后内容过短")
    return text


def _parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("未安装 pypdf，无法解析 PDF 文档") from exc

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        parts.append((page.extract_text() or "").strip())
    text = "\n".join(p for p in parts if p).strip()
    if len(text) < MIN_REQUIREMENT_CHARS:
        raise ValueError("PDF 文档解析后内容过短或无可提取文本")
    return text


def _parse_html(content: bytes, *, source_url: str | None = None) -> str:
    raw = _decode_text(content)
    lowered = raw.lower()
    if any(marker in lowered for marker in _LOGIN_MARKERS):
        if source_url and _parse_feishu_doc_url(source_url):
            raise ValueError(_feishu_login_required_message(has_open_api=False))
        raise ValueError("网页需要登录后才能访问，请改用可公开访问的链接，或上传本地文档 / 粘贴原文")

    raw = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", raw)
    raw = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", raw)
    raw = re.sub(r"(?is)<(br|p|div|li|h\d|tr)[^>]*>", "\n", raw)
    text = re.sub(r"<[^>]+>", " ", raw)
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text).strip()
    if len(text) < MIN_REQUIREMENT_CHARS:
        raise ValueError("网页内容过短，无法作为需求文档")
    return text
